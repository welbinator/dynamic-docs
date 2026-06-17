import uuid
from pathlib import Path
from flask import Blueprint, render_template, request, redirect, url_for, flash, current_app
from flask_login import login_required, current_user
from models import db, User, Organization, Document
from werkzeug.utils import secure_filename

settings_bp = Blueprint("settings", __name__)

PROVIDER_DEFAULTS = {
    "gemini": "gemini-2.5-flash",
    "openai": "gpt-4o",
    "anthropic": "claude-3-5-sonnet-20241022",
}

ALLOWED_EXTENSIONS = {"pdf", "docx", "doc", "txt", "md"}
MAX_FILE_BYTES = 50 * 1024 * 1024  # 50 MB


def admin_required(f):
    from functools import wraps
    @wraps(f)
    def decorated(*args, **kwargs):
        if not current_user.is_authenticated or not current_user.is_admin:
            return redirect(url_for("index"))
        return f(*args, **kwargs)
    return decorated


def uploads_dir() -> Path:
    d = Path(current_app.root_path) / "uploads"
    d.mkdir(exist_ok=True)
    return d


def extract_text(file_bytes: bytes, ext: str) -> str:
    """Extract plain text from uploaded file bytes."""
    if ext in ("txt", "md"):
        return file_bytes.decode("utf-8", errors="replace")

    if ext == "pdf":
        try:
            import pymupdf4llm
            import tempfile, os
            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
                tmp.write(file_bytes)
                tmp_path = tmp.name
            try:
                text = pymupdf4llm.to_markdown(tmp_path)
            finally:
                os.unlink(tmp_path)
            return text
        except ImportError:
            # Fallback: plain text extraction via pymupdf
            import fitz, io
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            return "\n\n".join(page.get_text() for page in doc)

    if ext in ("docx", "doc"):
        try:
            from docx import Document as DocxDocument
            import io
            docx_doc = DocxDocument(io.BytesIO(file_bytes))
            return "\n".join(p.text for p in docx_doc.paragraphs)
        except ImportError:
            return "[docx extraction requires python-docx — install it and re-upload]"

    return ""


@settings_bp.route("/settings", methods=["GET", "POST"])
@login_required
@admin_required
def settings_page():
    org = current_user.organization
    saved = False
    error = None

    if request.method == "POST":
        action = request.form.get("action")

        if action == "save_org":
            org_name = request.form.get("org_name", "").strip()
            if not org_name:
                error = "Organization name cannot be blank."
            else:
                org.name = org_name
                db.session.commit()
                saved = True

        elif action == "save_llm":
            provider = request.form.get("provider", "gemini")
            model = request.form.get("model", "").strip()
            api_key = request.form.get("api_key", "").strip()
            retrieval_mode = request.form.get("retrieval_mode", "complete")

            if provider not in PROVIDER_DEFAULTS:
                error = "Invalid provider selected."
            elif not model:
                error = "Model name cannot be blank."
            else:
                org.llm_provider = provider
                org.llm_model = model
                if api_key:
                    org.api_key = api_key
                org.retrieval_mode = retrieval_mode
                db.session.commit()
                saved = True

        elif action == "regenerate_invite":
            org.regenerate_invite_code()
            db.session.commit()
            saved = True

        elif action == "remove_user":
            user_id = request.form.get("user_id", type=int)
            if user_id:
                user = User.query.filter_by(id=user_id, org_id=org.id).first()
                if user and user.id != current_user.id:
                    db.session.delete(user)
                    db.session.commit()
                    saved = True

        elif action == "upload_doc":
            f = request.files.get("doc_file")
            if not f or not f.filename:
                error = "No file selected."
            else:
                ext = f.filename.rsplit(".", 1)[-1].lower() if "." in f.filename else ""
                if ext not in ALLOWED_EXTENSIONS:
                    error = f"Unsupported file type .{ext}. Allowed: {', '.join(sorted(ALLOWED_EXTENSIONS))}"
                else:
                    file_bytes = f.read()
                    if len(file_bytes) > MAX_FILE_BYTES:
                        error = "File exceeds 50 MB limit."
                    else:
                        stored_name = f"{uuid.uuid4().hex}.{ext}"
                        dest = uploads_dir() / stored_name
                        dest.write_bytes(file_bytes)

                        # Save the record immediately so the response returns fast.
                        # Text extraction (especially large PDFs) runs in a background thread.
                        doc = Document(
                            org_id=org.id,
                            original_name=f.filename,
                            stored_name=stored_name,
                            file_type=ext,
                            file_size=len(file_bytes),
                            extracted_text="[Processing… check back in a moment]",
                        )
                        db.session.add(doc)
                        db.session.commit()

                        import threading
                        _app = current_app._get_current_object()
                        _doc_id = doc.id
                        _bytes = file_bytes
                        _ext = ext

                        def _bg_extract():
                            with _app.app_context():
                                try:
                                    text = extract_text(_bytes, _ext)
                                except Exception as exc:
                                    text = f"[Extraction failed: {exc}]"
                                bg_doc = db.session.get(Document, _doc_id)
                                if bg_doc:
                                    bg_doc.extracted_text = text
                                    db.session.commit()

                        threading.Thread(target=_bg_extract, daemon=True).start()
                        saved = True

        elif action == "delete_doc":
            doc_id = request.form.get("doc_id", type=int)
            if doc_id:
                doc = Document.query.filter_by(id=doc_id, org_id=org.id).first()
                if doc:
                    stored = uploads_dir() / doc.stored_name
                    if stored.exists():
                        stored.unlink()
                    db.session.delete(doc)
                    db.session.commit()
                    saved = True

    members = User.query.filter_by(org_id=org.id).order_by(User.created_at).all()
    documents = Document.query.filter_by(org_id=org.id).order_by(Document.uploaded_at.desc()).all()
    invite_url = request.host_url.rstrip("/") + url_for("auth.join", invite_code=org.invite_code)

    return render_template(
        "settings.html",
        org=org,
        members=members,
        documents=documents,
        invite_url=invite_url,
        provider_defaults=PROVIDER_DEFAULTS,
        saved=saved,
        error=error,
    )
